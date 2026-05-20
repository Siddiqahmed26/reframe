"""Agent: generate a tailored cover letter.

Two responsibilities live here, kept together because they share a contract
(the LLM emits the plain-text body; the renderer turns that body into a
print-ready .docx that the frontend can offer as a download):

  - generate(jd, resume) -> str
        Calls the fast LLM to write the body text.
  - build_cover_letter_docx(text, candidate_name, contact, page_size) -> bytes
        Renders that text into an editorial-looking .docx in memory.
"""
from __future__ import annotations

import io
import logging
import re
import unicodedata
from datetime import date
from typing import Dict, Optional, Tuple

from backend.agents.llm import LLM
from backend.models import JDAnalysis, ResumeAnalysis


logger = logging.getLogger(__name__)


SYSTEM = """You write tailored, human-sounding cover letters. Your output is a single cover letter (no JSON, no preamble, no postamble), 220-340 words, three to four short paragraphs:

  - Opening: name the role and the company by name if known, plus one specific hook that ties the candidate's experience to a stated need in the JD.
  - Middle (1-2 paragraphs): two or three concrete achievements from the candidate's resume, mapped to the JD's responsibilities. Use real numbers and names from the resume.
  - Closing: a brief, confident line about wanting to discuss further, and a sign-off.

Rules:
  - Never invent experience. Only reference achievements that appear in the candidate's resume.
  - No clichés ("I am writing to express my interest..."). Open with substance.
  - No em dashes. Use commas, periods, or semicolons.
  - Plain text only. No markdown, no headers, no bullets.
  - Address it "Dear Hiring Manager," if no specific contact is mentioned.
  - End with "Sincerely," followed by the candidate's name from their resume.
"""


class CoverLetterIncompleteError(RuntimeError):
    """Raised when the LLM produced a body too short or syntactically
    unfinished. Callers should surface this as a retryable 503 rather than
    shipping a half-letter."""


def _looks_complete(text: str) -> bool:
    """A complete cover letter has at least 2 body paragraphs (after stripping
    greeting and sign-off) and the final body paragraph ends with sentence
    punctuation. Otherwise we treat it as truncated."""
    paragraphs = _split_body_paragraphs(text)
    # Drop the greeting (if any) and the sign-off (if any) — they don't count
    # toward "body" paragraphs.
    body = [
        p for p in paragraphs
        if not _GREETING_RE.match(p) and not _SIGNOFF_RE.match(p)
    ]
    if len(body) < 2:
        return False
    last = body[-1].rstrip()
    return last.endswith((".", "!", "?", '"'))


# Regexes shared with the renderers so completeness logic stays consistent.
_GREETING_RE = re.compile(r"^\s*(dear\s+|hello\s+|hi\s+|to\s+whom)", re.IGNORECASE)
_SIGNOFF_RE = re.compile(
    r"^\s*(sincerely|regards|best|yours|thank\s+you)[,\s]",
    re.IGNORECASE,
)


def generate(jd: JDAnalysis, resume: ResumeAnalysis, *, llm: "LLM | None" = None) -> str:
    user = f"""Job description analysis:
{jd.model_dump_json(indent=2)}

Candidate resume analysis:
{resume.model_dump_json(indent=2)}

Write the cover letter now.
"""
    # Default path uses the FAST model. Reasons:
    #   1. Cover-letter prose doesn't need the strongest model — it's
    #      single-shot generation with no schema to follow.
    #   2. The fast model on Groq (llama-3.1-8b-instant) has its own
    #      independent daily token quota (500K) separate from the main
    #      model's (gpt-oss-120b @ 200K). So the cover letter still
    #      generates even when the rewriter has exhausted the main quota.
    # BYOK injects a single user-provided model; respect it as-is.
    if llm is None:
        llm = LLM(fast=True)

    # First attempt at 2500 tokens — enough for a 340-word, 3-4 paragraph
    # editorial letter with full sentences and the typed sign-off.
    text = llm.complete(system=SYSTEM, user=user, max_tokens=2500, temperature=0.4)
    if _looks_complete(text):
        return text

    # Heuristic-based truncation retry. We can't easily plumb finish_reason
    # through the provider chain (different SDKs surface it differently), so
    # we trust the structural check: a body with <2 paragraphs OR ending
    # without sentence punctuation is almost certainly truncated. One retry
    # at 4000 tokens.
    logger.warning(
        "Cover letter looks truncated (paragraphs/punctuation check failed). "
        "Retrying once at max_tokens=4000."
    )
    retry = llm.complete(system=SYSTEM, user=user, max_tokens=4000, temperature=0.4)
    if _looks_complete(retry):
        return retry

    # Still bad after retry — surface as a controlled error. Better to fail
    # loudly than ship a half-letter the user might send to a recruiter.
    raise CoverLetterIncompleteError(
        "Cover letter generation incomplete: the model returned a body with "
        "fewer than 2 full paragraphs or an unfinished final sentence, even "
        "after retry. Please try again."
    )


# ── Editorial .docx renderer ──────────────────────────────────────────────
# Style palette
_NAVY = (26, 35, 57)         # header name
_SLATE = (110, 118, 138)     # contact strip
_LINK = (71, 92, 122)        # hyperlinks (slate-blue, restrained)
_BODY_INK = (38, 44, 58)     # body text
_BRAND_ACCENT = (91, 232, 197)  # the emerald-cyan rule

_HEADER_FONT = "Calibri"
_BODY_FONT = "Georgia"

# Type scale. Bumped from the spec defaults because users reported the
# previous output felt cramped — body 12pt at 1.4 line spacing fills a
# typical 280-340 word cover letter cleanly without looking dense.
_NAME_SIZE = 24
_CONTACT_SIZE = 10
_DATE_SIZE = 11.5
_BODY_SIZE = 12
_BODY_LINE = 1.4
_TYPED_NAME_SIZE = 12.5

# Dashes / curly quotes the LLM might still emit despite the system prompt.
# We never want them in the final document.
_TEXT_FALLBACKS = {
    "—": ", ",   # em dash → comma + space
    "–": "-",    # en dash → hyphen
    "‐": "-", "‑": "-", "‒": "-",
    "‘": "'", "’": "'",
    "“": '"', "”": '"',
    "…": "...",
}


def _sanitize_text(text: str) -> str:
    """Strip control characters and substitute typographic glyphs we don't
    want in a print document (em dashes, curly quotes)."""
    if not text:
        return ""
    out = []
    for ch in text:
        if ch in _TEXT_FALLBACKS:
            out.append(_TEXT_FALLBACKS[ch])
            continue
        cat = unicodedata.category(ch)
        # Drop control / format chars EXCEPT line break and tab.
        if cat.startswith("C") and ch not in ("\n", "\t"):
            continue
        out.append(ch)
    return "".join(out)


def _sanitize_name(name: Optional[str]) -> str:
    if not name:
        return ""
    cleaned = _sanitize_text(name).strip()
    # Collapse internal whitespace runs.
    return re.sub(r"\s+", " ", cleaned)


def _strip_handle(value: str, *, host_fragment: str, handle_prefix: str) -> str:
    """Render '<prefix>/<handle>' for a LinkedIn or GitHub field. The user's
    resume might already have the full URL, just the username, or the
    'in/username' / 'gh/username' shorthand. Normalize all of them."""
    v = (value or "").strip().lstrip("@").lstrip("/")
    if not v:
        return ""
    # Strip protocol + www
    v = re.sub(r"^https?://(www\.)?", "", v, flags=re.I)
    # Strip the host portion if present.
    low = v.lower()
    idx = low.find(host_fragment)
    if idx >= 0:
        v = v[idx + len(host_fragment):].lstrip("/")
    # The "in/" or "users/" path segment that LinkedIn/GitHub usernames sit
    # under should be dropped too (we'll re-add a short prefix).
    if v.lower().startswith("in/"):
        v = v[3:]
    return f"{handle_prefix}/{v}" if v else ""


def _full_url(value: str, *, host: str) -> str:
    """Return the canonical https URL for a LinkedIn or GitHub field. Tolerant
    of bare usernames, "in/user" shorthand, and full URLs (with or without
    protocol/www)."""
    v = (value or "").strip().lstrip("@").lstrip("/")
    if not v:
        return ""
    if v.lower().startswith(("http://", "https://")):
        return v
    if host.lower() in v.lower():
        return "https://" + v
    if host == "linkedin.com":
        path = v[3:] if v.lower().startswith("in/") else v
        return f"https://linkedin.com/in/{path.lstrip('/')}"
    return f"https://{host}/{v}"


def _format_contact_items(
    contact: Optional[Dict[str, str]],
) -> list[Tuple[str, Optional[str]]]:
    """Return ordered (display_text, optional_url) tuples for the contact strip.

    The URL is set for fields that should render as clickable hyperlinks:
    email (mailto:), linkedin, github. Plain fields (phone, location) carry
    None.  Renderers join these with a middle-dot separator.
    """
    if not contact:
        return []
    items: list[Tuple[str, Optional[str]]] = []

    email = (contact.get("email") or "").strip()
    if email:
        items.append((_sanitize_text(email), f"mailto:{email}"))

    phone = (contact.get("phone") or "").strip()
    if phone:
        items.append((_sanitize_text(phone), None))

    location = (contact.get("location") or "").strip()
    if location:
        items.append((_sanitize_text(location), None))

    linkedin_display = _strip_handle(
        contact.get("linkedin") or "",
        host_fragment="linkedin.com",
        handle_prefix="in",
    )
    if linkedin_display:
        url = _full_url(contact.get("linkedin") or "", host="linkedin.com")
        items.append((_sanitize_text(linkedin_display), url or None))

    github_display = _strip_handle(
        contact.get("github") or "",
        host_fragment="github.com",
        handle_prefix="gh",
    )
    if github_display:
        url = _full_url(contact.get("github") or "", host="github.com")
        items.append((_sanitize_text(github_display), url or None))

    return items


def _format_contact(contact: Optional[Dict[str, str]]) -> str:
    """Plain-text contact line (no hyperlinks) — used when we just need a
    string for testing or fallback rendering."""
    return " · ".join(text for text, _ in _format_contact_items(contact))


def _split_body_paragraphs(text: str) -> list[str]:
    """Split the LLM output into paragraphs. Prefer blank-line separators;
    fall back to single newlines if the model produced no blank lines."""
    cleaned = _sanitize_text(text or "").strip()
    if not cleaned:
        return []
    if "\n\n" in cleaned:
        chunks = [c.strip() for c in cleaned.split("\n\n")]
        return [c for c in chunks if c]
    # No blank-line breaks — treat each non-empty line as its own paragraph.
    return [ln.strip() for ln in cleaned.splitlines() if ln.strip()]


def build_cover_letter_docx(
    text: str,
    candidate_name: Optional[str] = None,
    contact: Optional[Dict[str, str]] = None,
    *,
    candidate_headline: Optional[str] = None,
    page_size: str = "letter",
) -> bytes:
    """Render the LLM cover-letter text into an editorial-looking .docx.

    Layout (top → bottom):
      1. Name in ALL CAPS, Calibri 22pt semi-bold navy, slightly expanded.
      2. Contact line in muted slate, separated by middle dots.
      3. Empty paragraph with a 0.75pt emerald-cyan bottom border (accent rule).
      4. Date in Georgia 11pt.
      5. Body paragraphs in Georgia 11pt, 1.35 line spacing, 0.25" first-line
         indent on continuation paragraphs.
      6. Sign-off: "Sincerely," + signature gap + typed name in Calibri 11.5pt.

    Returns the .docx bytes (seek(0) on a BytesIO).
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    from backend.writers.letter_styles import (
        set_run_font,
        set_paragraph_spacing,
        add_horizontal_rule,
        add_hyperlink_run,
    )

    name_clean = _sanitize_name(candidate_name)
    body_paragraphs = _split_body_paragraphs(text)
    if len(body_paragraphs) < 2:
        # Diagnostic: helps us tell apart parser bugs vs LLM truncation when
        # the output is suspiciously short. The renderer still emits what it
        # has — the agent layer is where we'd hard-fail.
        logger.warning(
            "Cover letter docx: only %d paragraph(s) detected after split; "
            "raw text len=%d.",
            len(body_paragraphs), len(text or ""),
        )

    # The LLM's prompt asks it to end with "Sincerely,\n{name}", which means
    # the typed name appears inside the body run. We render the typed name
    # ourselves in the header font, so strip it from the LLM tail to avoid
    # printing it twice.
    if name_clean and body_paragraphs:
        last = body_paragraphs[-1]
        trail_re = re.compile(
            r"\n\s*" + re.escape(name_clean) + r"\s*$",
            re.IGNORECASE,
        )
        body_paragraphs[-1] = trail_re.sub("", last).rstrip()

    doc = Document()

    # Page setup
    section = doc.sections[0]
    if page_size.lower() == "a4":
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    else:  # letter (default)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    # No header / footer content.
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True

    # 1. Candidate name — ALL CAPS, Calibri 22pt semi-bold, deep navy, slight
    #    letter-spacing for that "letterhead" feel. Left-aligned (editorial,
    #    not centered).
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(name_para, before_pt=0, after_pt=2)
    name_run = name_para.add_run((name_clean or "YOUR NAME").upper())
    # No manual letter-spacing — the PDF equivalent showed that even small
    # tracking dispersed the letters visually. The bold weight at 24pt
    # carries the header on its own.
    set_run_font(
        name_run,
        name=_HEADER_FONT,
        size_pt=_NAME_SIZE,
        bold=True,
        color_rgb=_NAVY,
    )

    # 1b. Optional headline — sits directly under the name, in a smaller
    #     muted slate so it reads as a tagline, not as the name itself.
    headline_clean = _sanitize_text(candidate_headline or "").strip()
    if headline_clean:
        headline_para = doc.add_paragraph()
        headline_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(headline_para, before_pt=0, after_pt=4)
        headline_run = headline_para.add_run(headline_clean)
        set_run_font(
            headline_run,
            name=_HEADER_FONT,
            size_pt=10.5,
            bold=False,
            color_rgb=_SLATE,
        )

    # 2. Contact strip — Calibri 10pt muted slate, single line, middle-dot
    #    separators. Items with a URL (email/linkedin/github) render as real
    #    hyperlinks (clickable in Word, underlined slate-blue).
    contact_items = _format_contact_items(contact)
    if contact_items:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(contact_para, before_pt=0, after_pt=6)
        for i, (text, url) in enumerate(contact_items):
            if i > 0:
                sep_run = contact_para.add_run(" · ")
                set_run_font(
                    sep_run,
                    name=_HEADER_FONT,
                    size_pt=_CONTACT_SIZE,
                    bold=False,
                    color_rgb=_SLATE,
                )
            if url:
                add_hyperlink_run(
                    contact_para,
                    url=url,
                    text=text,
                    name=_HEADER_FONT,
                    size_pt=_CONTACT_SIZE,
                    color_rgb=_LINK,
                )
            else:
                run = contact_para.add_run(text)
                set_run_font(
                    run,
                    name=_HEADER_FONT,
                    size_pt=_CONTACT_SIZE,
                    bold=False,
                    color_rgb=_SLATE,
                )

    # 3. Accent rule — empty paragraph carrying a bottom border.
    rule_para = doc.add_paragraph()
    set_paragraph_spacing(rule_para, before_pt=0, after_pt=8)
    add_horizontal_rule(
        rule_para,
        color_rgb=_BRAND_ACCENT,
        width_eighth_pt=6,   # 0.75pt
        space_pt=4,
    )

    # 4. Date — long form, locale-independent.
    date_para = doc.add_paragraph()
    set_paragraph_spacing(date_para, before_pt=10, after_pt=14)
    date_run = date_para.add_run(date.today().strftime("%B %d, %Y"))
    set_run_font(
        date_run,
        name=_BODY_FONT,
        size_pt=_DATE_SIZE,
        bold=False,
        color_rgb=_BODY_INK,
    )

    # 5. Body. The LLM emits the salutation, body paragraphs, and the
    #    sign-off line. We render every paragraph the LLM gave us, but
    #    sign-off-looking lines ("Sincerely,") get flush-left treatment
    #    and a signature gap below — body paragraphs get a small first-line
    #    indent for editorial polish.
    signoff_re = re.compile(r"^\s*(sincerely|regards|best|yours|thank\s+you)[,\s]", re.IGNORECASE)
    for i, para_text in enumerate(body_paragraphs):
        is_signoff_line = bool(signoff_re.match(para_text)) and len(para_text) < 40
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if i == 0:
            # Greeting / opening: no indent, breathing room before the next para.
            set_paragraph_spacing(p, before_pt=0, after_pt=12, line=_BODY_LINE)
        elif is_signoff_line:
            # Sign-off: flush left, generous gap before for the signature line.
            set_paragraph_spacing(p, before_pt=14, after_pt=24, line=_BODY_LINE)
        else:
            # Continuation body paragraph.
            set_paragraph_spacing(
                p,
                before_pt=0,
                after_pt=10,
                line=_BODY_LINE,
                first_line_indent_in=0.25,
            )
        run = p.add_run(para_text)
        set_run_font(
            run,
            name=_BODY_FONT,
            size_pt=_BODY_SIZE,
            bold=False,
            color_rgb=_BODY_INK,
        )

    # 6. Sign-off block.
    #    Detect whether the LLM already included a "Sincerely," line; if so,
    #    skip ours (avoid double sign-off). Otherwise render the canonical
    #    sign-off + signature gap + typed name.
    has_signoff = any(
        re.match(r"^\s*(sincerely|regards|best|yours)[,\s]", p, re.I)
        for p in body_paragraphs[-2:]
    )

    if not has_signoff:
        spacer = doc.add_paragraph()
        set_paragraph_spacing(spacer, before_pt=0, after_pt=10)

        signoff = doc.add_paragraph()
        set_paragraph_spacing(signoff, before_pt=0, after_pt=24, line=_BODY_LINE)
        signoff_run = signoff.add_run("Sincerely,")
        set_run_font(
            signoff_run,
            name=_BODY_FONT,
            size_pt=_BODY_SIZE,
            bold=False,
            color_rgb=_BODY_INK,
        )

    # Typed name — Calibri semi-bold for visual contrast with the serif body.
    if name_clean:
        typed = doc.add_paragraph()
        set_paragraph_spacing(typed, before_pt=0, after_pt=0)
        typed_run = typed.add_run(name_clean)
        set_run_font(
            typed_run,
            name=_HEADER_FONT,
            size_pt=_TYPED_NAME_SIZE,
            bold=True,
            color_rgb=_NAVY,
        )
    # If candidate_name is missing, omit the typed name entirely — a human
    # can fill it in. We deliberately don't render a "Your Name" placeholder
    # because it'd ship through email if the user didn't notice.

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def slugify_name(name: Optional[str]) -> str:
    """Produce a safe filename stem from a candidate name. Falls back to
    'candidate' when the input is empty or all non-word characters."""
    clean = _sanitize_name(name) or ""
    slug = re.sub(r"[^\w]+", "-", clean).strip("-").lower()
    return slug or "candidate"
