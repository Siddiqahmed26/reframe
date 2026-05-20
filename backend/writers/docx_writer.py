"""Format-preserving DOCX writer.

Strategy: we never create new paragraphs and never touch paragraph-level
properties. For each paragraph that has rewritten text, we:

  1. Take the new text string from the rewrite instructions.
  2. Place it inside the FIRST run of the paragraph, which carries the
     run-level formatting (font, size, bold/italic/underline, color).
  3. Clear the text of every other run in that paragraph, keeping their
     run/style elements intact so the XML structure does not change shape.

This means paragraph alignment, indentation, spacing, bullet markers, section
breaks, and page margins are entirely untouched. The visible difference
between the original and the rewritten document is the *text content* of
bullets and summary lines only.

Robustness: every RewriteInstruction is validated against the actual document
before its text is written. If the instruction's `original` field does NOT
match the text at `paragraph_index`, the rewrite is dropped with a warning.
This prevents the failure mode where the LLM emits a wrong index and our
output ends up with experience-section text dumped into the certifications
section.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Dict, List

from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from backend.models import RewriteInstruction
from backend.parsers.docx_parser import ParsedDocx


logger = logging.getLogger(__name__)


def _iter_all_paragraphs(doc: DocxDocument):
    """Yield every paragraph in body and in tables, in the same order the
    parser visited them. Returns tuples of (composite_index, paragraph)."""
    body_count = len(doc.paragraphs)
    for idx, p in enumerate(doc.paragraphs):
        yield idx, p
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    composite_idx = body_count + 1000 * t_idx + 100 * r_idx + 10 * c_idx + p_idx
                    yield composite_idx, p


def _replace_text_preserving_format(paragraph: Paragraph, new_text: str) -> None:
    """Put new_text into the paragraph while preserving every formatting
    attribute of the original first run."""
    runs = paragraph.runs
    if not runs:
        # An empty paragraph (rare in resumes for content lines). Just add a run.
        paragraph.add_run(new_text)
        return

    # First run gets the new text
    first = runs[0]
    first.text = new_text

    # Subsequent runs are emptied but their elements remain so style references
    # downstream (e.g. an italic run that styled a job title) don't disappear.
    for run in runs[1:]:
        run.text = ""


def _normalize_for_match(text: str) -> str:
    """Normalize a paragraph text for cross-check matching against the LLM's
    `original` field. Collapses whitespace, strips leading bullet glyphs,
    and lowercases — so a minor formatting difference between what we showed
    the LLM and what's literally in the doc doesn't trip the assertion."""
    if not text:
        return ""
    # Strip leading bullet markers + whitespace
    stripped = text.lstrip("•●▪◦■□▶►–—-· *\t ")
    # Collapse runs of whitespace
    import re
    return re.sub(r"\s+", " ", stripped).strip().lower()


def apply_rewrites(parsed: ParsedDocx, instructions: List[RewriteInstruction], out_path: str | Path) -> Path:
    """Apply all rewrite instructions to the original parsed document and save.

    Validates each instruction against the document before applying:
      1. The paragraph_index must exist in the document.
      2. The instruction's `original` text must match the text at that
         paragraph (modulo bullet glyphs and whitespace).
    Mismatches are logged at WARNING and SKIPPED — never written to a
    different paragraph silently. This is the guard against bullets ending
    up under the wrong section header.
    """
    # Index by paragraph index for O(1) lookup
    by_idx: Dict[int, RewriteInstruction] = {ins.paragraph_index: ins for ins in instructions}

    if not by_idx:
        # Nothing to rewrite; still save a copy so the API can return the file
        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        parsed.document.save(str(out))
        return out

    # Build a map composite_idx → paragraph so we can validate before writing.
    doc_paragraphs: Dict[int, Paragraph] = {
        idx: p for idx, p in _iter_all_paragraphs(parsed.document)
    }

    applied = 0
    skipped_missing = 0
    skipped_mismatch = 0

    for ins_idx, ins in by_idx.items():
        if not ins.rewritten or ins.rewritten == ins.original:
            continue
        p = doc_paragraphs.get(ins_idx)
        if p is None:
            skipped_missing += 1
            logger.warning(
                "Rewrite skipped — no paragraph at composite index %d. "
                "Instruction original starts: %r",
                ins_idx, (ins.original or "")[:60],
            )
            continue
        # Match check: the doc paragraph's text must look like the LLM's
        # claimed `original`. We allow bullet-glyph and whitespace drift.
        doc_text = _normalize_for_match(p.text)
        claim_text = _normalize_for_match(ins.original)
        # Tolerate truncation in either direction (LLM sometimes shortens
        # in its echo) by checking either is a prefix of the other.
        is_match = (
            doc_text == claim_text
            or (doc_text and claim_text and (
                doc_text.startswith(claim_text[:80]) or claim_text.startswith(doc_text[:80])
            ))
        )
        if not is_match:
            skipped_mismatch += 1
            logger.warning(
                "Rewrite skipped — paragraph at index %d does NOT match instruction's `original`. "
                "Doc says: %r | LLM claimed: %r",
                ins_idx, p.text[:60], (ins.original or "")[:60],
            )
            continue
        _replace_text_preserving_format(p, ins.rewritten)
        applied += 1

    logger.info(
        "DOCX writer: applied=%d, skipped_missing=%d, skipped_mismatch=%d",
        applied, skipped_missing, skipped_mismatch,
    )

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    parsed.document.save(str(out))
    return out
