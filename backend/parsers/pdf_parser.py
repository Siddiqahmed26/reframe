"""PDF parsing with span/bbox awareness so the PDF writer can edit text
in place.

We extract every text block on every page along with its bounding box, font,
size, color, and bold/italic flags. Each block becomes a ResumeParagraph whose
composite index encodes (page, block_idx) so the writer can locate it again.

The orchestrator treats ParsedPdf identically to ParsedDocx for the
LLM-facing pipeline (heuristic analysis, matching, rewriting). The dispatch
between writers happens by isinstance at the end.

Backwards compat: ``pdf_to_basic_docx`` and ``extract_text`` are retained so
older code paths and the MCP server keep working.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from backend.models import ResumeAnalysis, ResumeParagraph


class ResumeFormatError(ValueError):
    """Raised when the uploaded resume can't be processed for reasons the
    user can fix (image-only PDF, encrypted, corrupt). main.py maps this to
    HTTP 400 with the message; everything else falls through to a 500 with
    a logged traceback."""


SECTION_KEYWORDS = {
    "summary": ["summary", "profile", "objective", "about"],
    "experience": ["experience", "employment", "work history", "professional experience"],
    "education": ["education", "academic"],
    "skills": ["skills", "technical skills", "core competencies", "technologies"],
    "projects": ["projects", "selected projects", "personal projects"],
    "certifications": ["certifications", "certificates", "licenses"],
    "awards": ["awards", "honors", "achievements"],
    "publications": ["publications", "papers"],
}

BULLET_CHARS = ("•", "●", "·", "▪", "◦", "■", "□", "▶", "►", "–", "—", "-", "*")

# Encodes (page, block_idx, sub_idx) into a single int so it fits ResumeParagraph.index.
# sub_idx is the per-bullet split within a PyMuPDF text block (see _split_block_into_paragraphs).
_PAGE_STRIDE = 1_000_000
_BLOCK_STRIDE = 1_000


def _pack_index(page_idx: int, block_idx: int, sub_idx: int = 0) -> int:
    return page_idx * _PAGE_STRIDE + block_idx * _BLOCK_STRIDE + sub_idx


def _unpack_index(packed: int) -> tuple[int, int, int]:
    page, rem = divmod(packed, _PAGE_STRIDE)
    block, sub = divmod(rem, _BLOCK_STRIDE)
    return page, block, sub


@dataclass
class PdfSpan:
    """One contiguous run of text on a page, as PyMuPDF reports it."""
    text: str
    bbox: tuple[float, float, float, float]  # (x0, y0, x1, y1)
    font: str
    size: float
    color: int  # 0xRRGGBB
    flags: int  # PyMuPDF flag bitmask (bold/italic/etc.)


@dataclass
class PdfBlock:
    """One logical paragraph (often a bullet) split out of a PyMuPDF text block.

    PyMuPDF groups visually-clustered lines into a single block; for resumes
    this means an entire list of bullets under one job role becomes one block.
    The parser splits each block at bullet boundaries so each bullet becomes
    its own PdfBlock (with sub_idx > 0 for non-first splits). The writer then
    redacts and rewrites bullets individually.
    """
    page_idx: int
    block_idx: int
    sub_idx: int
    spans: List[PdfSpan]
    bbox: tuple[float, float, float, float]

    @property
    def text(self) -> str:
        return " ".join(s.text for s in self.spans if s.text).strip()


@dataclass
class ParsedPdf:
    """Holds the PyMuPDF document and a logical view of it. Mirrors ParsedDocx."""
    document: "object"  # pymupdf.Document; left untyped to keep import lazy
    blocks: List[PdfBlock]
    paragraphs: List[ResumeParagraph]
    full_text: str
    warnings: List[str] = field(default_factory=list)


def parse_pdf(path: str | Path) -> ParsedPdf:
    """Parse a PDF into spans + a paragraph view. Raises ValueError if the PDF
    is image-only (no extractable text), since we can't rewrite it."""
    import pymupdf  # lazy: keeps module import cheap and offline-friendly

    p = Path(path)
    doc = pymupdf.open(str(p))

    blocks: List[PdfBlock] = []
    paragraphs: List[ResumeParagraph] = []
    current_section: Optional[str] = None

    for page_idx, page in enumerate(doc):
        # "dict" mode gives us blocks -> lines -> spans with full formatting
        page_dict = page.get_text("dict")
        for block_idx, block in enumerate(page_dict.get("blocks", [])):
            if block.get("type") != 0:  # 0 = text block, 1 = image
                continue

            # Split this PyMuPDF block into one PdfBlock per logical bullet
            # so the rewriter can edit each bullet individually.
            for sub_idx, (sub_spans, sub_bbox) in enumerate(
                _split_block_into_paragraphs(block)
            ):
                blk = PdfBlock(
                    page_idx=page_idx,
                    block_idx=block_idx,
                    sub_idx=sub_idx,
                    spans=sub_spans,
                    bbox=sub_bbox,
                )
                blocks.append(blk)
                packed = _pack_index(page_idx, block_idx, sub_idx)
                rp, current_section = _classify(blk.text, packed, current_section)
                paragraphs.append(rp)

    full_text = "\n".join(p.text for p in paragraphs if p.text)

    warnings: List[str] = []
    if not full_text.strip():
        # No extractable text at all -> image-based PDF
        doc.close()
        raise ResumeFormatError(
            "This PDF appears to be image-based (no extractable text). "
            "Please upload a text-based PDF or a .docx file."
        )

    return ParsedPdf(
        document=doc,
        blocks=blocks,
        paragraphs=paragraphs,
        full_text=full_text,
        warnings=warnings,
    )


def _line_starts_bullet(line: dict) -> bool:
    """True if the line's leading non-whitespace character is a bullet glyph."""
    for span in line.get("spans", []):
        text = (span.get("text") or "").lstrip()
        if not text:
            continue
        return text[0] in BULLET_CHARS
    return False


def _line_x0(line: dict):
    """Leftmost x-coordinate of the line's first non-empty span, or None."""
    for span in line.get("spans", []):
        if (span.get("text") or "").strip():
            return float(span.get("bbox", (0,))[0])
    return None


def _merge_bbox(a, b):
    return (min(a[0], b[0]), min(a[1], b[1]), max(a[2], b[2]), max(a[3], b[3]))


def _spans_from_line(line: dict) -> List[PdfSpan]:
    out: List[PdfSpan] = []
    for span in line.get("spans", []):
        text = span.get("text", "")
        if not text:
            continue
        out.append(PdfSpan(
            text=text,
            bbox=tuple(span.get("bbox", (0, 0, 0, 0))),
            font=span.get("font", ""),
            size=float(span.get("size", 11.0)),
            color=int(span.get("color", 0)),
            flags=int(span.get("flags", 0)),
        ))
    return out


def _split_block_into_paragraphs(block: dict):
    """Yield (spans, bbox) pairs by splitting a PyMuPDF block at bullet
    starts. Lines that don't begin with a bullet glyph are treated as
    continuations of the previous group (e.g. wrapped lines of one bullet,
    or a header line followed by bullets).

    If the block contains no bullets at all, it yields a single group
    containing all lines combined — same behavior as before.
    """
    lines = block.get("lines", [])
    if not lines:
        return

    has_any_bullet = any(_line_starts_bullet(ln) for ln in lines)

    if not has_any_bullet:
        # No bullets to split on. Emit the whole block as one paragraph.
        all_spans: List[PdfSpan] = []
        bbox = None
        for ln in lines:
            ln_spans = _spans_from_line(ln)
            if not ln_spans:
                continue
            all_spans.extend(ln_spans)
            ln_bbox = tuple(ln.get("bbox", (0, 0, 0, 0)))
            bbox = ln_bbox if bbox is None else _merge_bbox(bbox, ln_bbox)
        if all_spans:
            yield all_spans, (bbox or (0, 0, 0, 0))
        return

    # Split rules:
    #   (a) a bullet line always starts a new group
    #   (b) a non-bullet line AFTER a bullet line is one of two things:
    #       - continuation of the wrapping bullet (indented under the bullet
    #         text, so its x0 is well to the RIGHT of the bullet glyph's x0)
    #       - a role/section header for the NEXT bullet group (x0 at or
    #         near the bullet glyph's x0, i.e. NOT indented)
    #       We split iff the non-bullet line is NOT indented relative to
    #       the most recent bullet glyph. Tolerance: 4pt.
    #
    # This is the key fix that prevents the writer from redacting role
    # headers (HulkHire/HP/Flipkart etc.) when rewriting the bullet above.

    INDENT_TOLERANCE = 4.0  # points

    current_spans: List[PdfSpan] = []
    current_bbox = None
    current_has_bullet = False
    last_bullet_x0 = None

    for ln in lines:
        ln_spans = _spans_from_line(ln)
        if not ln_spans:
            continue
        ln_bbox = tuple(ln.get("bbox", (0, 0, 0, 0)))
        line_is_bullet = _line_starts_bullet(ln)
        line_x0 = _line_x0(ln)

        should_split = False
        if line_is_bullet and current_spans:
            should_split = True
        elif (
            not line_is_bullet
            and current_has_bullet
            and last_bullet_x0 is not None
            and line_x0 is not None
            and line_x0 <= last_bullet_x0 + INDENT_TOLERANCE
        ):
            # Non-bullet, not indented under the bullet = it's a header
            # for the next group, not a continuation. Split here.
            should_split = True

        if should_split and current_spans:
            yield current_spans, (current_bbox or (0, 0, 0, 0))
            current_spans = []
            current_bbox = None
            current_has_bullet = False

        current_spans.extend(ln_spans)
        current_bbox = ln_bbox if current_bbox is None else _merge_bbox(current_bbox, ln_bbox)
        if line_is_bullet:
            current_has_bullet = True
            last_bullet_x0 = line_x0

    if current_spans:
        yield current_spans, (current_bbox or (0, 0, 0, 0))


def _classify(text: str, idx: int, current_section: Optional[str]) -> tuple[ResumeParagraph, Optional[str]]:
    text = text.strip()
    is_heading = False
    is_bullet = False

    if text and len(text) <= 40:
        letters = sum(1 for c in text if c.isalpha())
        if letters:
            upper_ratio = sum(1 for c in text if c.isupper()) / letters
            if upper_ratio > 0.7 and len(text.split()) <= 5:
                is_heading = True

    # Section detection: walk each line within this paragraph and look for a
    # section keyword. Necessary because PyMuPDF often groups "WORK EXPERIENCE"
    # together with the first job's header line, making the combined text > 40
    # chars (so a single-text check would miss it).
    section = current_section
    for line in (l.strip() for l in text.split("\n") if l.strip()):
        if len(line) > 40:
            continue
        lowered_line = line.lower()
        for key, kws in SECTION_KEYWORDS.items():
            if any(kw in lowered_line for kw in kws):
                section = key
                if line == text:
                    is_heading = True
                break
        if section != current_section:
            break

    if text.startswith(BULLET_CHARS):
        is_bullet = True

    return (
        ResumeParagraph(
            index=idx,
            text=text,
            style=None,
            is_heading=is_heading,
            section=section,
            is_bullet=is_bullet,
        ),
        section,
    )


def parsed_pdf_to_analysis(parsed: ParsedPdf) -> ResumeAnalysis:
    """Heuristic ResumeAnalysis from a ParsedPdf. Mirrors parsed_to_analysis
    in docx_parser so the orchestrator doesn't have to special-case PDFs."""
    import re

    analysis = ResumeAnalysis(paragraphs=parsed.paragraphs)

    for p in parsed.paragraphs:
        if p.text:
            analysis.candidate_name = p.text
            break

    head_text = "\n".join(p.text for p in parsed.paragraphs[:8])
    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", head_text)
    phone_match = re.search(r"(\+?\d[\d\s().-]{7,}\d)", head_text)
    linkedin_match = re.search(r"(linkedin\.com/in/[\w-]+)", head_text, re.IGNORECASE)
    github_match = re.search(r"(github\.com/[\w-]+)", head_text, re.IGNORECASE)

    if email_match:
        analysis.contact["email"] = email_match.group(0)
    if phone_match:
        analysis.contact["phone"] = phone_match.group(0).strip()
    if linkedin_match:
        analysis.contact["linkedin"] = linkedin_match.group(0)
    if github_match:
        analysis.contact["github"] = github_match.group(0)

    sections_seen = set()
    for p in parsed.paragraphs:
        if not p.text or p.is_heading:
            continue
        sec = p.section
        if sec:
            sections_seen.add(sec)
        if sec == "summary":
            analysis.summary = (analysis.summary + " " + p.text).strip() if analysis.summary else p.text
        elif sec == "experience" and p.is_bullet:
            analysis.experience_bullets.append(p.text)
        elif sec == "education":
            analysis.education.append(p.text)
        elif sec == "projects" and p.is_bullet:
            analysis.projects.append(p.text)
        elif sec == "skills":
            for token in re.split(r"[,;|/]", p.text):
                token = token.strip()
                if token and len(token) <= 40:
                    analysis.skills.append(token)
    analysis.sections_found = sorted(sections_seen)
    return analysis


# ---------------------------------------------------------------------------
# Legacy helpers kept for backwards compatibility with callers that haven't
# migrated to parse_pdf yet (notably the old DOCX fallback path).
# ---------------------------------------------------------------------------

def extract_text(path: str | Path) -> str:
    """Pull all text out of a PDF as plain string. Kept for callers that
    want a quick text dump."""
    try:
        import pdfplumber
        p = Path(path)
        lines: List[str] = []
        with pdfplumber.open(str(p)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    stripped = line.strip()
                    if stripped:
                        lines.append(stripped)
        return "\n".join(lines)
    except Exception:
        # Fall back to PyMuPDF if pdfplumber isn't available
        import pymupdf
        doc = pymupdf.open(str(path))
        out = "\n".join(page.get_text("text") for page in doc)
        doc.close()
        return out


def pdf_to_basic_docx(pdf_path: str | Path, out_path: str | Path) -> Path:
    """Legacy: convert PDF to a flat .docx. The new pipeline prefers parse_pdf,
    but this is kept as a fallback (e.g. for the cover-letter-only endpoint)."""
    from docx import Document
    from docx.shared import Pt

    text = extract_text(pdf_path)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for line in text.splitlines():
        para = doc.add_paragraph(line)
        if line.startswith(BULLET_CHARS):
            para.style = doc.styles["List Bullet"]
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
