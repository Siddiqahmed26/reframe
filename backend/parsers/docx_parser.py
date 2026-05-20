"""Parse .docx into a logical representation while keeping a reference to the
original document so the writer can put rewritten text back in the same place
without disturbing formatting."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from backend.models import ResumeParagraph, ResumeAnalysis


# Section headers we look for. Case-insensitive substring match.
SECTION_KEYWORDS = {
    "summary": ["summary", "profile", "objective", "about"],
    "experience": ["experience", "employment", "work history", "professional experience"],
    "education": ["education", "academic"],
    "skills": ["skills", "technical skills", "core competencies", "technologies"],
    "projects": ["projects", "selected projects", "personal projects"],
    "certifications": ["certifications", "certificates", "licenses"],
    "awards": ["awards", "honors", "achievements"],
    "publications": ["publications", "papers"],
}


@dataclass
class ParsedDocx:
    """Holds both the python-docx Document and a logical view of it."""
    document: DocxDocument
    paragraphs: List[ResumeParagraph]
    full_text: str


def _classify_paragraph(p: Paragraph, idx: int, current_section: Optional[str]) -> Tuple[ResumeParagraph, Optional[str]]:
    """Return a ResumeParagraph and the (possibly updated) current section."""
    text = p.text.strip()
    style_name = p.style.name if p.style is not None else None

    # Detect heading
    is_heading = False
    if style_name and "Heading" in style_name:
        is_heading = True
    # Heuristic: short ALL CAPS or Title Case lines are often section headers
    if text and len(text) <= 40:
        upper_ratio = sum(1 for c in text if c.isupper()) / max(1, sum(1 for c in text if c.isalpha()))
        if upper_ratio > 0.7 and len(text.split()) <= 5:
            is_heading = True

    # Detect section by keyword match
    section = current_section
    if is_heading or (text and len(text) <= 40):
        lowered = text.lower()
        for key, kws in SECTION_KEYWORDS.items():
            if any(kw in lowered for kw in kws):
                section = key
                is_heading = True
                break

    # Bullet detection: docx bullets usually live in list paragraph styles,
    # but plenty of resumes just use a leading dash/dot character.
    is_bullet = False
    if style_name and ("List" in style_name or "Bullet" in style_name):
        is_bullet = True
    elif text.startswith(("•", "●", "·", "-", "*", "▪", "◦")):
        is_bullet = True

    return (
        ResumeParagraph(
            index=idx,
            text=text,
            style=style_name,
            is_heading=is_heading,
            section=section,
            is_bullet=is_bullet,
        ),
        section,
    )


def parse_docx(path: str | Path) -> ParsedDocx:
    """Parse a .docx file and return both the original Document and a logical view."""
    doc = Document(str(path))
    paragraphs: List[ResumeParagraph] = []
    current_section: Optional[str] = None

    # Walk paragraphs in body
    for idx, para in enumerate(doc.paragraphs):
        rp, current_section = _classify_paragraph(para, idx, current_section)
        paragraphs.append(rp)

    # Also walk tables (many resumes have a left sidebar in a table)
    table_offset = len(doc.paragraphs)
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    composite_idx = table_offset + 1000 * t_idx + 100 * r_idx + 10 * c_idx + p_idx
                    rp, current_section = _classify_paragraph(para, composite_idx, current_section)
                    paragraphs.append(rp)

    # Smart line-join for `full_text`: certificate/education paragraphs are
    # sometimes split by a soft line break that python-docx surfaces as two
    # separate paragraphs. We merge them in the analysis text so the matcher
    # and cover-letter agents see complete titles. The underlying
    # `paragraphs` list (and its indices) is NOT mutated — the writer still
    # needs stable indices for in-place rewriting.
    full_text = _build_joined_full_text(paragraphs)

    return ParsedDocx(document=doc, paragraphs=paragraphs, full_text=full_text)


# Words that look like a paragraph was split mid-sentence by a soft line break.
_CONJUNCTION_TAILS = ("and", "&", "or", "of", "the", "with", "in", "for", "to", "a", "an")


def _build_joined_full_text(paragraphs: List[ResumeParagraph]) -> str:
    """Concatenate paragraph texts with newlines, but merge consecutive
    paragraphs in the same section when the first ends with a conjunction
    (and/or/of/the/&/...) or a comma AND the second starts with a lowercase
    letter — the typical signature of an accidental soft-line-break split."""
    lines: list[str] = []
    prev_p: Optional[ResumeParagraph] = None
    for p in paragraphs:
        if not p.text:
            prev_p = p
            continue
        if (
            lines
            and prev_p is not None
            and prev_p.section == p.section
            and _looks_like_split(prev_p.text, p.text)
        ):
            # Merge into the previous emitted line.
            lines[-1] = lines[-1].rstrip(" ,;:") + " " + p.text.lstrip()
        else:
            lines.append(p.text)
        prev_p = p
    return "\n".join(lines)


def _looks_like_split(prev_text: str, next_text: str) -> bool:
    """Return True if `prev_text` looks like it was cut mid-clause and
    `next_text` continues it."""
    p = prev_text.rstrip()
    if not p or not next_text:
        return False
    # Ends with a sentence terminator → not a split.
    if p[-1] in ".!?":
        return False
    last_token = p.split()[-1].lower().rstrip(",;:")
    ends_with_conjunction = (
        last_token in _CONJUNCTION_TAILS
        or p[-1] in ",;&"
    )
    if not ends_with_conjunction:
        return False
    # Next paragraph should start with a lowercase letter — uppercase would
    # be a new sentence / heading, not a continuation.
    first_char = next_text.lstrip()[:1]
    return first_char.islower()


def _balance_phone(raw: str) -> str:
    """Tidy a phone number we just matched: balance parentheses, collapse
    internal runs of whitespace, drop trailing punctuation.

    Common breakage: the regex captures "+91) 9663908548" because the resume
    has "(+91)" but the opening paren was outside the match window. We
    detect orphaned parens and either close or strip them so we never ship
    a stray ")" in the rendered contact strip.
    """
    s = (raw or "").strip()
    if not s:
        return ""
    opens = s.count("(")
    closes = s.count(")")
    if opens > closes:
        # Missing close paren — drop the open(s) entirely; the surrounding
        # whitespace is usually fine on its own.
        s = s.replace("(", "", opens - closes)
    elif closes > opens:
        # Orphan close paren. If the front of the string looks like
        # "+CC " with a country code we know was meant to be parenthesized,
        # restore the open paren. Otherwise strip the trailing close.
        import re
        m = re.match(r"\+(\d{1,3})\b", s)
        if m and ")" in s and "(" not in s:
            cc = m.group(1)
            # Replace "+CC" with "(+CC)" at the start.
            s = s.replace("+" + cc, "(+" + cc, 1)
            # Leave the existing ")" in place; now parens balance.
        else:
            s = s.replace(")", "", closes - opens)
    # Collapse any whitespace runs.
    import re
    s = re.sub(r"\s+", " ", s).strip(" ,;:.-")
    return s


def parsed_to_analysis(parsed: ParsedDocx) -> ResumeAnalysis:
    """Build a ResumeAnalysis from a ParsedDocx using lightweight heuristics."""
    analysis = ResumeAnalysis(paragraphs=parsed.paragraphs)
    sections_seen = set()

    # First non-empty line is usually the candidate name
    for p in parsed.paragraphs:
        if p.text:
            analysis.candidate_name = p.text
            break

    # Simple contact extraction from the first ~8 paragraphs.
    import re
    head_text = "\n".join(p.text for p in parsed.paragraphs[:8])

    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", head_text)

    # Phone: handle plain, "(+91) 9663...", "+91-9663-...", and "(123) 456-7890"
    # forms. The previous regex required the first char to be `\+?\d`, so a
    # leading `(` got skipped and we captured "+91) 9663..." with a stray
    # closing paren. New regex accepts an optional leading `(` and the rest
    # of the number, then we balance any orphan parens in post-processing.
    phone_match = re.search(
        r"(\+?\(?\d{1,3}\)?[\s\-.]?\d[\d\s\-.()]{6,}\d)",
        head_text,
    )

    # LinkedIn / GitHub: accept either a full URL ("linkedin.com/in/user"),
    # a path-only form ("in/user", "gh/user"), or a bare handle on a line
    # tagged with a host name.
    linkedin_match = re.search(
        r"linkedin\.com/(?:in/)?([\w.\-]+)",
        head_text,
        re.IGNORECASE,
    )
    github_match = re.search(
        r"github\.com/([\w.\-]+)",
        head_text,
        re.IGNORECASE,
    )

    # Location: typical resumes put "City, State" or "City, Country" near
    # the contact line. Match a comma-separated pair of capitalized tokens
    # NOT containing digits or '@' (filters out address lines that have
    # numbers).
    location_match = re.search(
        r"\b([A-Z][a-zA-Z]+(?:\s[A-Z][a-zA-Z]+)?,\s[A-Z][a-zA-Z]+(?:\s[A-Z][a-zA-Z]+)?)\b",
        head_text,
    )

    if email_match:
        analysis.contact["email"] = email_match.group(0)
    if phone_match:
        analysis.contact["phone"] = _balance_phone(phone_match.group(1).strip())
    if linkedin_match:
        analysis.contact["linkedin"] = "linkedin.com/in/" + linkedin_match.group(1)
    if github_match:
        analysis.contact["github"] = "github.com/" + github_match.group(1)
    if location_match:
        analysis.contact["location"] = location_match.group(1).strip()

    # Bucket paragraphs by section
    for p in parsed.paragraphs:
        if not p.text or p.is_heading:
            continue
        sec = p.section
        if sec:
            sections_seen.add(sec)
        if sec == "summary":
            if analysis.summary:
                analysis.summary += " " + p.text
            else:
                analysis.summary = p.text
        elif sec == "experience" and p.is_bullet:
            analysis.experience_bullets.append(p.text)
        elif sec == "education":
            analysis.education.append(p.text)
        elif sec == "projects" and p.is_bullet:
            analysis.projects.append(p.text)
        elif sec == "skills":
            # Skills are often comma- or pipe-separated
            for token in re.split(r"[,;|/]", p.text):
                token = token.strip()
                if token and len(token) <= 40:
                    analysis.skills.append(token)

    analysis.sections_found = sorted(sections_seen)
    return analysis
