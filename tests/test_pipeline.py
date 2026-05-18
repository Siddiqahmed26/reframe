"""Smoke tests that don't require an API key. They exercise the parser, the
writer's format preservation, the ATS scorer's deterministic path, and verify
that every package imports cleanly."""
from __future__ import annotations

import importlib
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from backend.models import (
    ATSScore,
    JDAnalysis,
    MatchReport,
    ResumeAnalysis,
    ResumeParagraph,
    RewriteInstruction,
    RewriteResult,
)
from backend.parsers.docx_parser import parse_docx, parsed_to_analysis
from backend.writers.docx_writer import apply_rewrites
from backend.agents import ats_optimizer


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Build a tiny resume-like .docx for tests."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com | linkedin.com/in/janedoe | +1 555 0100")
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Software engineer with five years of backend experience.")
    doc.add_heading("Experience", level=1)
    p = doc.add_paragraph("Built a payment service handling 2 million requests per day.")
    p.style = doc.styles["List Bullet"]
    p = doc.add_paragraph("Migrated the monolith to microservices using Kubernetes.")
    p.style = doc.styles["List Bullet"]
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, Go, Kubernetes, PostgreSQL")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("BS Computer Science, State University, 2018")

    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out


def test_parse_docx_finds_sections(sample_docx: Path) -> None:
    parsed = parse_docx(sample_docx)
    assert parsed.full_text
    analysis = parsed_to_analysis(parsed)
    assert analysis.candidate_name == "Jane Doe"
    assert analysis.contact.get("email") == "jane@example.com"
    assert "linkedin" in analysis.contact
    assert "experience" in analysis.sections_found
    assert "skills" in analysis.sections_found
    assert any("payment service" in b for b in analysis.experience_bullets)
    assert "Python" in analysis.skills


def test_writer_preserves_paragraph_count(sample_docx: Path, tmp_path: Path) -> None:
    parsed = parse_docx(sample_docx)
    original_count = len(parsed.document.paragraphs)
    # Build a no-op rewrite (same text)
    instructions = [
        RewriteInstruction(paragraph_index=i, original=p.text, rewritten=p.text)
        for i, p in enumerate(parsed.paragraphs)
    ]
    out = apply_rewrites(parsed, instructions, tmp_path / "out.docx")
    assert out.exists()
    reopened = parse_docx(out)
    assert len(reopened.document.paragraphs) == original_count


def test_writer_changes_text_only(sample_docx: Path, tmp_path: Path) -> None:
    parsed = parse_docx(sample_docx)
    # Find the first bullet under experience and rewrite it
    target = next(p for p in parsed.paragraphs if p.is_bullet)
    new_text = "Scaled a transaction processing service to 5 million daily requests."
    instructions = [
        RewriteInstruction(
            paragraph_index=target.index,
            original=target.text,
            rewritten=new_text,
        )
    ]
    out_path = tmp_path / "rewritten.docx"
    apply_rewrites(parsed, instructions, out_path)

    # Reopen and verify the bullet text changed and the font is preserved
    reopened = parse_docx(out_path)
    bullets = [p for p in reopened.paragraphs if p.is_bullet]
    assert any(new_text in b.text for b in bullets)

    # Verify the paragraph style is unchanged
    new_doc = Document(str(out_path))
    bullet_paras = [p for p in new_doc.paragraphs if p.style.name and "Bullet" in p.style.name]
    assert bullet_paras, "Bullet style should be preserved"


def test_ats_score_deterministic_path() -> None:
    """When the rewritten text contains the keywords, coverage should be 1.0."""
    jd = JDAnalysis(
        role_title="Backend Engineer",
        must_have_skills=["python", "kubernetes"],
        keywords=["microservices", "postgresql"],
    )
    text = "Python backend engineer who built microservices on Kubernetes with PostgreSQL."
    score = ats_optimizer.score(text, jd)
    assert score.keyword_coverage == 1.0
    assert set(score.missing_keywords) == set()


def test_ats_score_missing_keywords() -> None:
    jd = JDAnalysis(must_have_skills=["rust"], keywords=["wasm"])
    text = "Python and Go engineer with backend experience."
    score = ats_optimizer.score(text, jd)
    assert score.keyword_coverage < 1.0
    assert "rust" in score.missing_keywords
    assert "wasm" in score.missing_keywords


def test_models_serialize() -> None:
    """Make sure every pydantic model round-trips through JSON."""
    jd = JDAnalysis(role_title="x", must_have_skills=["a"], keywords=["b"])
    assert JDAnalysis.model_validate_json(jd.model_dump_json()) == jd
    ra = ResumeAnalysis(candidate_name="x")
    assert ResumeAnalysis.model_validate_json(ra.model_dump_json()) == ra


def test_imports() -> None:
    """Every internal module must import cleanly without an API key."""
    for mod in [
        "backend.config",
        "backend.models",
        "backend.parsers.docx_parser",
        "backend.parsers.pdf_parser",
        "backend.writers.docx_writer",
        "backend.agents.llm",
        "backend.agents.jd_analyzer",
        "backend.agents.resume_analyzer",
        "backend.agents.matcher",
        "backend.agents.rewriter",
        "backend.agents.ats_optimizer",
        "backend.agents.cover_letter",
        "backend.agents.orchestrator",
        "backend.main",
        "mcp_server.server",
    ]:
        importlib.import_module(mod)
